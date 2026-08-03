from docx import Document

doc = Document(r'C:\Users\emman\Desktop\descarga1\NanoRuntime_Paper.docx')
full = '\n'.join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

print('=== VERIFICACION DE COMPLETITUD ===')
print('')

# 1. Structure
sections = ['Abstract','1. Introduction','2. Related Work','3. System Design',
            '4. Evaluation','5. Discussion','6. Conclusion','Acknowledgments','References']
print('1. ESTRUCTURA:')
all_ok = True
for s in sections:
    ok = s in full
    if not ok: all_ok = False
    print('  ' + ('[OK]' if ok else '[MISSING]') + ' ' + s)

# 2. Key content
checks = [
    ('NanoRuntime','Nombre del sistema'),
    ('Emmanuel Higuita','Autor'),
    ('Resource-Aware Graceful Degradation','Contribucion 1'),
    ('26.7% reduction','Metrica clave RAM'),
    ('< 1 MB','Metrica varianza RSS'),
    ('OPPO CPH2557','Dispositivo 1'),
    ('Samsung Galaxy A30s','Dispositivo 2'),
    ('50/50','Exito OPPO'),
    ('30/30','Exito Samsung'),
    ('+296 MB','RAM net OPPO'),
    ('+351 MB','RAM net Samsung'),
    ('0.43 tok/s','Throughput 7B'),
    ('2.90 tok/s','Throughput 1.5B OPPO'),
    ('10.74','PC ablation tok/s'),
    ('1840','PC ablation RSS NR'),
    ('MMLU','Benchmark MMLU'),
    ('HumanEval','Benchmark HumanEval'),
    ('90.0%','MMLU score'),
    ('66.7%','HumanEval score'),
    ('Shapiro-Wilk','Test estadistico'),
    ('Mann-Whitney','Test estadistico'),
    ('bootstrap','Metodo CI'),
    ('psutil','Herramienta medicion'),
    ('adb shell','Metodo Android'),
    ('/proc/meminfo','Fuente RAM'),
    ('Rust','Lenguaje'),
    ('llama.cpp','Backend'),
    ('Q4_K_M','Cuantizacion'),
    ('DeepSeek-R1','Modelo 7B'),
    ('Shannon','Entropia'),
    ('Graceful Degradation','Mecanismo'),
    ('KV cache','Concepto clave'),
    ('Rappi','Afiliacion'),
]

print('\n2. CONTENIDO CLAVE:')
for pattern, desc in checks:
    ok = pattern in full
    if not ok: all_ok = False
    print('  ' + ('[OK]' if ok else '[MISSING]') + ' ' + desc)

# 3. Tables
print('\n3. TABLAS:')
tables = doc.tables
print('  Total: ' + str(len(tables)))
for i, t in enumerate(tables):
    if t.rows and t.rows[0].cells:
        first = t.rows[0].cells[0].text[:50]
        rows = len(t.rows)
        print('  Tabla ' + str(i+1) + ': ' + str(rows) + ' filas - ' + first)

# 4. Stats
words = full.split()
print('\n4. ESTADISTICAS:')
print('  Palabras: ' + str(len(words)))
print('  Paragrafos: ' + str(len(doc.paragraphs)))

# 5. Final verdict
print('\n' + ('=' * 50))
if all_ok:
    print('  VEREDICTO: DOCUMENTO COMPLETO')
    print('  Sin secciones faltantes. Sin datos faltantes.')
else:
    print('  VEREDICTO: FALTAN ELEMENTOS')
print('=' * 50)
