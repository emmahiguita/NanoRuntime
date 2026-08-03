from docx import Document
import re

doc = Document(r'C:\Users\emman\Desktop\descarga1\NanoRuntime_Paper.docx')
full_text = '\n'.join([p.text for p in doc.paragraphs])
issues = []

# AI-trace phrases
ai_phrases = ['delve','dive deep','furthermore','moreover','in conclusion',
              'as we have seen','tapestry','landscape','realm','robust',
              'leverage','crucially','notably','it is worth noting',
              'it is important to']
for phrase in ai_phrases:
    if phrase.lower() in full_text.lower():
        issues.append('AI-trace phrase: ' + phrase)

# Grammar errors
if re.search(r'\ba\s+11', full_text):
    issues.append('Grammar: "a 11" should be "an 11"')

# Double spaces
if '  ' in full_text:
    # Count them
    count = len(re.findall(r'  ', full_text))
    issues.append('Double spaces found: ' + str(count) + ' occurrences')

# Check sections exist
sections = ['Abstract','1. Introduction','2. Related Work','3. System Design',
            '4. Evaluation','5. Discussion','6. Conclusion','Acknowledgments','References']
missing = [s for s in sections if s not in full_text]
if missing:
    issues.append('Missing sections: ' + ', '.join(missing))

# Check for Spanish words in English text
spanish_words = ['ató','ión','ción','más','sólo','éste','cómo','cuál']
for w in spanish_words:
    if w in full_text:
        issues.append('Spanish word found: ' + w)

if not issues:
    print('NO ISSUES FOUND - Document is clean')
else:
    print('ISSUES FOUND:')
    for i in issues:
        print('  - ' + i)

print('')
print('Total paragraphs:', len(doc.paragraphs))
print('Total chars:', len(full_text))
for s in sections:
    print('  Section [' + s + ']: ' + ('OK' if s in full_text else 'MISSING'))
