#!/usr/bin/env python3
"""
_docx_builder.py — Shared DOCX document builder with DI.
Single source of truth for page setup, styles, and helper methods.
Previously duplicated across export_docx.py and export_docx_v2.py (~70 lines).

Usage:
    from _docx_builder import DocxBuilder
    builder = DocxBuilder()
    builder.add_title("Paper Title")
    builder.add_heading("Section 1", level=1)
    builder.add_para("Body text...")
    builder.save("output.docx")
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class DocxBuilder:
    """Encapsulates Document creation, page setup, and content building.
    Inject via constructor. No global state — each instance is independent."""

    def __init__(self) -> None:
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    # ── Page and style setup ────────────────────────────────────────
    def _setup_page(self) -> None:
        for section in self.doc.sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

    def _setup_styles(self) -> None:
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.15

    # ── Content helpers ─────────────────────────────────────────────
    def add_title(self, text: str, size: int = 16) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"

    def add_author(self, text: str, size: int = 12) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"

    def add_heading(self, text: str, level: int | None = None) -> None:
        """Add a section heading. If level is None, auto-detects:
        - Text starting with 'X.Y' (e.g. '3.1') → level 2 (sub-section)
        - Anything else → level 1 (section)."""
        import re
        if level is None:
            level = 2 if re.match(r"^\d+\.\d+\b", text) else 1
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.name = "Times New Roman"
        if level == 1:
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(18)
        else:
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(12)

    def add_para(
        self, text: str, bold: bool = False, italic: bool = False, size: int = 11
    ) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"

    def add_code(self, text: str) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Cm(1)

    def add_table(self, headers: list[str], rows: list[list[str]]) -> None:
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(9)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
        self.doc.add_paragraph()

    # ── Output ──────────────────────────────────────────────────────
    def save(self, path: str | Path) -> None:
        self.doc.save(str(path))
